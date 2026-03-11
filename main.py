import os
try:
    import truststore
    truststore.inject_into_ssl()
except ImportError:
    pass

from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from rich.console import Console
from rich.prompt import Prompt
from models.schemas import ResearchInput, ResearchProposal
from agents.supervisor import SupervisorAgent
from config import settings

console = Console()


def collect_input() -> ResearchInput:
    console.print("\n[bold cyan]=== AI Co-Scientist for Engineering ===[/bold cyan]\n")
    domain = Prompt.ask("[white]연구 도메인[/white]", default="광학 센서")
    objective = Prompt.ask("[white]연구 목표[/white]")
    constraints_raw = Prompt.ask("[white]제약조건 (쉼표로 구분)[/white]", default="저전력, 소형화")
    constraints = [c.strip() for c in constraints_raw.split(",")]
    background = Prompt.ask("[white]배경 정보 (없으면 Enter)[/white]", default="")
    return ResearchInput(
        domain=domain,
        objective=objective,
        constraints=constraints,
        background=background or None
    )


def _add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        run.font.size = Pt(13)
    return p


def _add_markdown_body(doc: Document, text: str):
    """마크다운 텍스트를 Word 단락으로 변환"""
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
        elif stripped.startswith("### "):
            _add_heading(doc, stripped[4:], level=2)
        elif stripped.startswith("## "):
            _add_heading(doc, stripped[3:], level=2)
        elif stripped.startswith("# "):
            _add_heading(doc, stripped[2:], level=2)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("```"):
            pass
        else:
            doc.add_paragraph(stripped)


def save_proposal(proposal: ResearchProposal, output_dir: str = None) -> tuple[Path, Path]:
    """마크다운(.md)과 Word(.docx) 두 형식으로 저장. (md_path, docx_path) 반환"""
    output_dir = Path(output_dir or settings.output_dir)
    output_dir.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    stem = f"{timestamp}_{proposal.input.domain.replace(' ', '_')}"

    # ── 1. 마크다운 저장 ──────────────────────────────
    md_path = output_dir / f"{stem}.md"
    md_content = f"""# {proposal.title}

**생성일시:** {proposal.created_at}
**도메인:** {proposal.input.domain}
**목표:** {proposal.input.objective}
**제약조건:** {', '.join(proposal.input.constraints)}

---

## 선택된 접근법

**{proposal.selected_idea.title}**

{proposal.selected_idea.description}

접근법: {proposal.selected_idea.approach}

---

## 설계 사양

{proposal.design_spec}

---

## 실험 계획

{proposal.experiment_plan}

---

## 시뮬레이션 방법

{proposal.simulation_suggestion}

---

## 참고문헌

{chr(10).join(proposal.references)}
"""
    md_path.write_text(md_content, encoding="utf-8")

    # ── 2. Word 문서 저장 ──────────────────────────────
    docx_path = output_dir / f"{stem}.docx"
    doc = Document()

    # 제목
    title_p = doc.add_heading(proposal.title, level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title_p.runs:
        title_p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        title_p.runs[0].font.size = Pt(20)

    # 메타 정보
    meta = doc.add_paragraph()
    meta.add_run(f"생성일시: {proposal.created_at}\n").font.size = Pt(10)
    meta.add_run(f"도메인: {proposal.input.domain}  |  ").font.size = Pt(10)
    meta.add_run(f"목표: {proposal.input.objective}  |  ").font.size = Pt(10)
    meta.add_run(f"제약조건: {', '.join(proposal.input.constraints)}").font.size = Pt(10)
    doc.add_paragraph()

    # 선택된 접근법
    _add_heading(doc, "선택된 접근법", level=1)
    p = doc.add_paragraph()
    p.add_run(proposal.selected_idea.title).bold = True
    doc.add_paragraph(proposal.selected_idea.description)
    doc.add_paragraph(f"접근법: {proposal.selected_idea.approach}")
    doc.add_paragraph()

    # 설계 사양
    _add_heading(doc, "설계 사양", level=1)
    _add_markdown_body(doc, proposal.design_spec)
    doc.add_paragraph()

    # 실험 계획
    _add_heading(doc, "실험 계획", level=1)
    _add_markdown_body(doc, proposal.experiment_plan)
    doc.add_paragraph()

    # 시뮬레이션 방법
    _add_heading(doc, "시뮬레이션 방법", level=1)
    _add_markdown_body(doc, proposal.simulation_suggestion)
    doc.add_paragraph()

    # 참고문헌
    _add_heading(doc, "참고문헌", level=1)
    for ref in proposal.references:
        doc.add_paragraph(ref, style="List Number")

    doc.save(str(docx_path))
    return md_path, docx_path


def main():
    research_input = collect_input()
    supervisor = SupervisorAgent()
    proposals = supervisor.run(research_input)
    console.print(f"\n[bold green]제안서 {len(proposals)}개 저장 완료![/bold green]")
    for proposal in proposals:
        md_path, docx_path = save_proposal(proposal)
        console.print(f"\n  [cyan]{proposal.selected_idea.title}[/cyan]")
        console.print(f"    Markdown : {md_path}")
        console.print(f"    Word     : {docx_path}")


if __name__ == "__main__":
    main()
